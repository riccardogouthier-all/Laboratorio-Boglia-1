from docx import Document
from docx.shared import RGBColor, Inches
from datetime import datetime
from pathlib import Path

from student_stats import (
    distribuzione_voti,
    peggiori_per_materia,
    studenti_a_rischio,
)
from student_charts import genera_tutti_i_grafici


SEZIONI_DISPONIBILI: dict[int, str] = {
    1: "Riepilogo validazione",
    2: "Indicatori di classe (medie per materia)",
    3: "Classifiche (top 5 e peggiori per materia)",
    4: "Analisi rischi (media < 6 o assenze > 15)",
    5: "Esiti finali (promossi/bocciati)",
    6: "Fasce di rendimento",
    7: "Classifiche ordinate (bubble/quick sort)",
}


def chiedi_sezioni() -> set[int]:
    """
    Mostra il menu interattivo e restituisce il set di numeri-sezione
    scelti dall'utente.

    Ritorna un set[int] con i numeri delle sezioni da includere,
    oppure l'insieme completo {1..7} se l'utente digita 'all'.
    """
    print("\n SEZIONI DISPONIBILI PER REPORT:")
    for num, descrizione in SEZIONI_DISPONIBILI.items():
        print(f"  {num}. {descrizione}")
    print("\n Inserisci i numeri delle sezioni separati da virgola (es: 1,2,5):")
    print("  'all' o NESSUN TESTO per includerle tutte\n")

    while True:
        risposta = input(" > ").strip().lower()
    # elif comando in ("generate-validate", "gv"):

        if risposta in ("all", ""):
            return set(SEZIONI_DISPONIBILI.keys())

        try:
            scelte = {int(s.strip()) for s in risposta.split(",") if s.strip()}
        except ValueError:
            print(" [ERRORE] Inserisci solo numeri separati da virgola, oppure 'all'.")
            continue

        non_validi = scelte - set(SEZIONI_DISPONIBILI.keys())
        if non_validi:
            print(f" [ERRORE] Numeri non validi: {sorted(non_validi)}. Scegli tra 1 e 7.")
            continue

        if not scelte:
            print(" [ERRORE] Devi selezionare almeno una sezione.")
            continue

        return scelte


def genera_report_docx(
    config: dict,
    validi: list[dict],
    scartati: list[dict],
    stats: dict,
    top5: list[dict],
    fasce: dict | None = None,
    classifica_media: list[dict] | None = None,
    classifica_assenze: list[dict] | None = None,
    genera_grafici: bool = True,
    sezioni: set[int] | None = None,
) -> Path:
    """
    Genera il report .docx in report/.

    *sezioni* è un set[int] (1-7) che indica quali sezioni includere.
    Se None o vuoto vengono incluse tutte.  Usa :func:`chiedi_sezioni`
    per la selezione interattiva.
    """
    # ── normalizzazione sezioni ───────────────────────────────────────────────
    if not sezioni:
        sezioni = set(SEZIONI_DISPONIBILI.keys())   # default: tutte

    includi = sezioni.__contains__   # alias corto per leggibilità

    oggi     = datetime.now()
    oradiadesso = datetime.now().strftime("%H%M%S")
    percorso = Path("report") / (f"report_{oggi.strftime('%Y%m%d')}" + "_" + f"{oradiadesso}.docx")
    percorso.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # ── helpers di formattazione ──────────────────────────────────────────────

    def aggiungi_titolo(testo: str, livello: int = 1):
        doc.add_heading(testo, level=livello)

    def aggiungi_riga(testo: str, grassetto: bool = False, colore: RGBColor | None = None):
        """Aggiunge un paragrafo normale, opzionalmente in grassetto o colorato."""
        p = doc.add_paragraph()
        run = p.add_run(testo)
        run.bold = grassetto
        if colore:
            run.font.color.rgb = colore

    def aggiungi_voce_lista(testo: str, stile: str = "List Bullet"):
        doc.add_paragraph(testo, style=stile)

    # ── INTESTAZIONE ─────────────────────────────────────────────────────────
    aggiungi_titolo("Report Studenti — Pipeline Python", livello=0)

    # ── SEZIONE 1 — Riepilogo validazione ────────────────────────────────────
    if includi(1):
        aggiungi_riga(f"Classe:              {config['classe']}", grassetto=True)
        aggiungi_riga(f"Data generazione:    {oggi.strftime('%d/%m/%Y %H:%M:%S')}")
        aggiungi_riga(f"Studenti totali:     {len(validi) + len(scartati)}")
        aggiungi_riga(f"Studenti validi:     {len(validi)}")
        aggiungi_riga(f"Studenti scartati:   {len(scartati)}")

    # ── SEZIONE 2 — Indicatori di classe (medie per materia) ─────────────────
    if includi(2):
        aggiungi_titolo("Statistiche per Materia")
        for materia, s in stats.items():
            aggiungi_riga(materia, grassetto=True)
            aggiungi_voce_lista(f"Media:            {s['media']}")
            aggiungi_voce_lista(f"Mediana:          {s['mediana']}")
            aggiungi_voce_lista(f"Dev. standard:    {s['stdev']}")
            aggiungi_voce_lista(f"Voto minimo:      {s['min']}")
            aggiungi_voce_lista(f"Voto massimo:     {s['max']}")

    # ── SEZIONE 3 — Classifiche ───────────────────────────────────────────────
    if includi(3):
        aggiungi_titolo("Distribuzione Voti (globale)")
        dist        = distribuzione_voti(studenti=validi)
        totale_voti = sum(dist.values())
        for voto, conteggio in dist.items():
            pct = round(conteggio / totale_voti * 100, 1) if totale_voti else 0
            aggiungi_riga(f"  Voto {voto:>2}:  {conteggio:>3} studenti  ({pct:>4}%)")

        aggiungi_titolo("Peggiori per Materia (top 3)")
        for materia, lista in peggiori_per_materia(studenti=validi, top_n=3).items():
            aggiungi_riga(materia, grassetto=True)
            for pos, s in enumerate(lista, 1):
                aggiungi_voce_lista(
                    f"{pos}. {s['nome']} {s['cognome']}  —  voto: {s['voto']}"
                )

        aggiungi_titolo("Top 5 Studenti")
        for i, s in enumerate(top5, 1):
            aggiungi_voce_lista(
                f"{i}. {s['nome']} {s['cognome']}  —  "
                f"media: {s.get('media', s.get('media_personale', '—'))}  "
                f"assenze: {s['assenze']}"
            )

    # ── SEZIONE 4 — Analisi rischi ────────────────────────────────────────────
    if includi(4):
        rischio = studenti_a_rischio(
            studenti=validi,
            soglia_voto=float(config.get("soglia_voto", 6.0)),
            soglia_assenze=int(config.get("soglia_assenze", 15)),
        )
        aggiungi_titolo("Studenti a Rischio")
        aggiungi_riga(
            f"Media insufficiente (<{config.get('soglia_voto', 6.0)}):  "
            f"{len(rischio['media_insufficiente'])} studenti"
        )
        aggiungi_riga(
            f"Troppe assenze (>{config.get('soglia_assenze', 15)}):       "
            f"{len(rischio['troppe_assenze'])} studenti"
        )
        aggiungi_riga(
            f"Entrambi i problemi:          {len(rischio['entrambi'])} studenti"
        )
        if rischio["entrambi"]:
            aggiungi_riga("Dettaglio (entrambi i problemi):", grassetto=True)
            for s in rischio["entrambi"]:
                aggiungi_voce_lista(
                    f"⚠  {s['nome']} {s['cognome']}  —  "
                    f"media: {s['media']}   assenze: {s['assenze']}",
                    stile="List Bullet",
                )

    # ── SEZIONE 5 — Esiti finali ──────────────────────────────────────────────
    ha_esito = validi and "esito" in validi[0]
    if includi(5) and ha_esito:
        n_promossi = sum(1 for s in validi if s["esito"] == "Promosso")
        n_bocciati = len(validi) - n_promossi
        aggiungi_titolo("Esiti Finali")
        aggiungi_riga(f"Promossi:   {n_promossi}")
        aggiungi_riga(f"Bocciati:   {n_bocciati}")
        bocciati = [s for s in validi if s["esito"] == "Bocciato"]
        if bocciati:
            aggiungi_riga("Dettaglio bocciati:", grassetto=True)
            for s in bocciati:
                motivi_str = "; ".join(s.get("motivi_esito", []))
                aggiungi_voce_lista(
                    f"• {s['nome']} {s['cognome']}  —  "
                    f"voto finale: {s.get('voto_finale', '—')}  "
                    f"debiti: {len(s.get('debiti', []))}  "
                    f"[{motivi_str}]"
                )

    # ── SEZIONE 6 — Fasce di rendimento ──────────────────────────────────────
    if includi(6) and fasce:
        aggiungi_titolo("Fasce di Rendimento")
        for nome_fascia, lista_fascia in fasce.items():
            aggiungi_riga(f"{nome_fascia}  —  {len(lista_fascia)} studenti", grassetto=True)
            for s in lista_fascia:
                debiti_str = ", ".join(s.get("debiti", [])) or "nessuno"
                aggiungi_voce_lista(
                    f"{s['nome']} {s['cognome']}  —  "
                    f"media: {s.get('media_personale', '—')}  "
                    f"voto: {s.get('voto_finale', '—')}  "
                    f"debiti: {debiti_str}"
                )

    # ── SEZIONE 7 — Classifiche ordinate ─────────────────────────────────────
    if includi(7):
        if classifica_media:
            aggiungi_titolo("Classifica per Media (bubble sort)")
            for pos, s in enumerate(classifica_media, 1):
                aggiungi_voce_lista(
                    f"{pos:>3}. {s['nome']} {s['cognome']}  —  "
                    f"media: {s.get('media_personale', '—')}  "
                    f"esito: {s.get('esito', '—')}"
                )

        if classifica_assenze:
            aggiungi_titolo("Classifica per Assenze (quick sort, crescente)")
            for pos, s in enumerate(classifica_assenze, 1):
                aggiungi_voce_lista(
                    f"{pos:>3}. {s['nome']} {s['cognome']}  —  "
                    f"assenze: {s['assenze']}  "
                    f"media: {s.get('media_personale', '—')}"
                )

    # ── STEP 17 — GENERAZIONE GRAFICI (integrato nel report) ─────────────────
    if genera_grafici and validi and ha_esito:
        aggiungi_titolo("Grafici e Visualizzazioni")
        try:
            grafici = genera_tutti_i_grafici(validi, config)

            # Titoli leggibili per ogni grafico, nello stesso ordine di generazione
            titoli_grafici = {
                "istogrammi": "Distribuzione voti per materia",
                "barre_comparative": "Confronto media voti per materia",
                "scatter_correlazione": "Correlazione tra materie",
                "pie_chart": "Esiti finali (promossi vs bocciati)",
                "heatmap": "Heatmap voti per studente",
                "media_vs_assenze": "Media personale vs assenze",
            }

            n_inseriti = 0
            for chiave, percorso_grafico in grafici.items():
                if not percorso_grafico:
                    continue
                aggiungi_riga(titoli_grafici.get(chiave, chiave), grassetto=True)
                doc.add_picture(str(percorso_grafico), width=Inches(6))
                n_inseriti += 1

            aggiungi_riga(f"Grafici generati con successo in report/charts/ ({n_inseriti} file)")
            print("[Step 17] Grafici generati con successo in report/charts/")
        except Exception as e:
            aggiungi_riga(f"⚠  Avvertimento: impossibile generare grafici: {e}")
            print(f"[Step 17] ⚠  Avvertimento: impossibile generare grafici: {e}")

    # ── SALVATAGGIO ───────────────────────────────────────────────────────────
    doc.save(percorso)
    print(f"[Step 8] Report DOCX salvato in: {percorso}")
    return percorso