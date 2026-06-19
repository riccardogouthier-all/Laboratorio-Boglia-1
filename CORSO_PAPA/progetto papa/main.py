"""
student_pipeline — main.py
Pipeline completa per generare, validare e analizzare dati studenti.
Uso:
python main.py generate   → crea il CSV con studenti casuali
python main.py validate   → valida il CSV e produce il JSON
python main.py report     → calcola statistiche e produce il report
python main.py all        → esegue tutto in sequenza
"""

import sys           # Step 10 — CLI
import json          # Step 1, 5, 6 — config e JSON
import csv           # Step 3, 4 — lettura/scrittura CSV
import re            # Step 4 — validazione con espressioni regolari
import shutil        # Step 9 — copia file (backup)
from pathlib  import Path      # Step 0, 3 — gestione percorsi
from datetime import datetime #, date # usato per generare studenti hardcoded //  # Step 2, 3, 8 — date e timestamp
from collections import Counter      # Step 4, 7 — conteggio errori

from docx import Document

# import dai file py del progetto
from installer_dependencies import install_procedure

from database_set import (
    carica_configdb,
    main as database_main)

from database_reader import (leggi_studenti_da_db)

from student_logic import (
    calcola_media_e_voto_finale,
    assegna_debiti_formativi,
    determina_esito,
    raggruppa_per_fascia_rendimento,
    bubble_sort_studenti,
    quick_sort_studenti,
)

from student_stats import (
    media_per_materia,       # sostituisce calcola_statistica  (Step 6)
    classifica_studenti,     # sostituisce classifica_studenti (Step 7)
    peggiori_per_materia,    # nuovo — peggiori N studenti per ogni materia
    distribuzione_voti,      # nuovo — quanti 6, 7, 8… per materia o globale
    studenti_a_rischio)       # nuovo — media < soglia o assenze > soglia

from student_charts import genera_tutti_i_grafici

from genera_report import genera_report_docx

def crea_cartelle():            # STEP 0 - Preparazione: creazione delle cartelle di progetto
    """Crea la struttura di cartelle del progetto se non esistono."""
    path = Path.cwd()#("PROGETTO-Student Analytics Pipeline")      #"CORSO_BOGLIA","basi di programmazione",
    base_path = path

    cartelle = [
        base_path / "data/input",
        base_path / "data/output",
        base_path / "data/backup",
        base_path / "report",
        base_path / "report/charts",  # ← AGGIUNTO per i grafici
        base_path / "configurations"
    ]
    for cartella in cartelle:
        cartella.mkdir(parents=True, exist_ok=True)
    print("[Step 0] Cartelle create/verificate.")

def carica_config() -> dict:       # STEP 1 - dizionario impostazioni
    """Crea config.json con valori di default se non esiste, poi lo legge."""
    path = Path.cwd()                       #("PROGETTO-Student Analytics Pipeline")    #"CORSO_BOGLIA","basi di programmazione",
    CONFIG_PATH = path /"configurations"/ "config.json"
    

    DEFAULT_CONFIG = {
        "numero_studenti": 50,
        "voto_min": 2,
        "voto_max": 10,
        # "materie": ["Matematica", "Informatica", "Italiano"],
        "materie": ["Matematica", "Informatica", "Italiano", "Storia", "Inglese", "Educazione Fisica"],
        "classe": "5A",
        "pesi_materie": {"Matematica": 1, "Informatica": 1, "Italiano": 1, "Storia": 1, "Inglese": 1, "Educazione Fisica": 1}
        }

    if not CONFIG_PATH.exists():
        with open(CONFIG_PATH, "w", encoding="utf-8") as f:
            json.dump(DEFAULT_CONFIG, f, indent=2, ensure_ascii=False)
        print("[Step 1] config.json creato con valori di default.")
    else:
        print("[Step 1] config.json trovato, caricamento in corso.")

    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        config = json.load(f)
    return config       # dizionario impostazioni

def salva_su_csv(studenti: list[dict], config:dict) -> Path:       # STEP 3 - SALVATAGGIO SU CSV - path sistema x la lista python
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    path = Path("data","input")    
    percorso = path / f"studenti_{timestamp}.csv"

    materie = config["materie"]
    nomicampi = ["id","nome","cognome","data_nascita","email"] + materie +["assenze"]

    with open(percorso, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames= nomicampi)
        writer.writeheader()

        for studente in studenti:
            riga = {
                "id": studente["id"],
                "nome": studente["nome"],
                "cognome": studente["cognome"],
                "data_nascita": studente["data_nascita"],
                "email": studente["email"],
                "assenze": studente["assenze"]
            }
            for materia in materie:
                riga[materia] = studente["voti"][materia]
            writer.writerow(riga)

    print(f"[Step 3] CSV salvato in: {percorso}")
    return percorso       # path sistema x la lista python

def aggiungi_errori(percorso_csv: Path, config: dict) -> None:       # STEP 2B - Aggiunta studenti non validi per attivare controllo errori
    with open(percorso_csv, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames
        righe = list(reader)
    
    ultimo_id = max(int(r["id"]) for r in righe) if righe else 0

    studenti_da_scartare = [
            {
            "id":            ultimo_id + 1,
            "nome":          "Carlo",
            "cognome":       "Verdi",
            "data_nascita":  "99-99-9999",        # data non valida
            "email":         "carlo.verdiATscuola.it",  # email senza @
            **{m: config["voto_max"] + 3 for m in config["materie"]},  # voti fuori range
            "assenze":       5,
            },
            {
            "id":            ultimo_id + 2,
            "nome":          "Paola",
            "cognome":       "Neri",
            "data_nascita":  "2005-13-40",        # mese e giorno impossibili
            "email":         "paola.neri@scuola.it",  # email valida
            **{m: config["voto_max"] - config["voto_max"] for m in config["materie"]},  # voti = 0
            "assenze":       3,
            }
            ]
    
    righe.extend(studenti_da_scartare)

    with open(percorso_csv, "w", newline="", encoding="utf-8") as f:
        writer = csv.DictWriter(f, fieldnames= fieldnames)
        writer.writeheader()
        writer.writerows(righe)
    print(f"[Step 2B] studenti non validi aggiunti in: {percorso_csv}")

def valida_studenti(percorso_csv: Path, config: dict) -> tuple[list,list,Counter]:       # STEP 4 - Lettura CSV e validazione - ritorna lista validi e scartati python e conteggio errori (contatore)
    # Pattern regex per validare l'email: testo@testo.dominio
    PATTERN_EMAIL = re.compile(r"^[\w.\-]+@[\w.\-]+\.\w{2,}$")
    # Pattern per la data in formato YYYY-MM-DD
    PATTERN_DATA  = re.compile(r"^\d{4}-\d{2}-\d{2}$")

    validi = []
    scartati = []
    conteggio_errori = Counter()
    materie = config["materie"]

    with open(percorso_csv, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for riga in reader:
            errori = []

            if not PATTERN_EMAIL.match(riga["email"]):
                errori.append("email_non_valida")
                conteggio_errori["email_non_valida"] += 1

            if not PATTERN_DATA.match(riga["data_nascita"]):
                errori.append("data_nascita_non_valida")
                conteggio_errori["data_nascita_non_valida"] += 1
            else:
                try:
                    datetime.strptime(riga["data_nascita"], "%Y-%m-%d")
                except ValueError:
                    errori.append("data_nascita_non_valida")
                    conteggio_errori["data_nascita_non_valida"] += 1

            for materia in materie:
                try:
                    voto = int(riga[materia])
                    if not (config["voto_min"] <= voto <= config["voto_max"]):
                        errori.append("voto_fuori_range")
                        conteggio_errori["voto_fuori_range"] += 1
                except (ValueError, KeyError):
                        errori.append("voto_non_numerico")
                        conteggio_errori["voto_non_numerico"] += 1

            studente = {
                "id":           int(riga["id"]),
                "nome":         riga["nome"],
                "cognome":      riga["cognome"],
                "data_nascita": riga["data_nascita"],
                "email":        riga["email"],
                "assenze":      int(riga.get("assenze", 0)),
                "voti":         {m: int(riga[m]) for m in materie if riga.get(m, "").isdigit()},
                # Extra B — timestamp di importazione
                "timestamp_import": datetime.now().isoformat(),
                }
            
            if errori:
                studente["errori"] = errori
                scartati.append(studente)
            else:
                validi.append(studente)

    print(f"[Step 4] Validi: {len(validi)} | Scartati: {len(scartati)}")
    print(f"         Errori trovati: {dict(conteggio_errori) or 'nessuno'}")
    return validi, scartati, conteggio_errori       # ritorna lista validi, scartati python e conteggio errori (contatore)

def salva_json_validi(studenti: list[dict]) -> Path:            # STEP 5 - Conversione PERCORSO CSV → JSON - ritorna il percorso json_validi
    path = Path("data/output")       #("CORSO_BOGLIA","basi di programmazione","PROGETTO-Student Analytics Pipeline") 
    percorso = path / "studenti_validi.json"
    with open(percorso, "w", encoding= "utf-8") as f:
        json.dump(studenti, f, indent=2, ensure_ascii=False)
    print(f"[Step 5] JSON validi salvati in: {percorso}")
    return percorso         # percorso json_validi

def salva_json_scartati(scartati: list[dict]) -> Path:          # STEP 5 - Conversione PERCORSO CSV → JSON - ritorno il percorso json_scartati
    path = Path("data/output")       #("CORSO_BOGLIA","basi di programmazione","PROGETTO-Student Analytics Pipeline") 
    percorso = path / "studenti_scartati.json"
    with open(percorso, "w", encoding= "utf-8") as f:
        json.dump(scartati, f, indent=2, ensure_ascii=False)
    print(f"[Step 5] JSON scartati salvati in: {percorso} | ({len(scartati)} record)")
    return percorso         # percorso json_scartati




'''
def genera_report(config: dict, validi: list[dict], scartati: list[dict], stats: dict, top5: list[dict], fasce: dict | None = None, classifica_media: list[dict] | None = None, classifica_assenze: list[dict] | None = None, genera_grafici: bool = True) -> Path:
    # |stats| da student_stats.media_per_materia   → {materia: {media, mediana, stdev, min, max, soglia_floor}}
    # |top5| da student_stats.classifica_studenti → lista dal campo "migliori"
    """Genera il file di testo del report completo in report/."""
    oggi     = datetime.now()
    percorso = Path("report") / f"report_{oggi.strftime('%Y%m%d')}.txt"
    percorso.parent.mkdir(parents=True, exist_ok=True)

    righe = []
    righe.append("=" * 60)
    righe.append("  REPORT STUDENTI — PIPELINE PYTHON")
    righe.append("=" * 60)
    righe.append(f"Classe:              {config['classe']}")
    righe.append(f"Data generazione:    {oggi.strftime('%d/%m/%Y %H:%M:%S')}")
    righe.append(f"Studenti totali:     {len(validi) + len(scartati)}")
    righe.append(f"Studenti validi:     {len(validi)}")
    righe.append(f"Studenti scartati:   {len(scartati)}")

    # ── STATISTICHE PER MATERIA  (student_stats.media_per_materia) ───────────────
    righe.append("")
    righe.append("-" * 60)
    righe.append("  STATISTICHE PER MATERIA")
    righe.append("-" * 60)
    for materia, s in stats.items():
        righe.append(f"\n  {materia}")
        righe.append(f"    Media:            {s['media']}")
        righe.append(f"    Mediana:          {s['mediana']}")
        righe.append(f"    Dev. standard:    {s['stdev']}")
        righe.append(f"    Voto minimo:      {s['min']}")
        righe.append(f"    Voto massimo:     {s['max']}")

    # ── DISTRIBUZIONE VOTI  (student_stats.distribuzione_voti) ───────────────────
    righe.append("")
    righe.append("-" * 60)
    righe.append("  DISTRIBUZIONE VOTI (globale)")
    righe.append("-" * 60)
    dist = distribuzione_voti(studenti=validi)
    totale_voti = sum(dist.values())
    for voto, conteggio in dist.items():
        barra = "█" * conteggio
        pct   = round(conteggio / totale_voti * 100, 1) if totale_voti else 0
        righe.append(f"  {voto:>2}: {barra:<52} {conteggio:>3}  ({pct:>4}%)")

    # ── PEGGIORI PER MATERIA  (student_stats.peggiori_per_materia) ───────────────
    righe.append("")
    righe.append("-" * 60)
    righe.append("  PEGGIORI PER MATERIA (top 3)")
    righe.append("-" * 60)
    for materia, lista in peggiori_per_materia(studenti=validi, top_n=3).items():
        righe.append(f"\n  {materia}")
        for pos, s in enumerate(lista, 1):
            righe.append(f"    {pos}. {s['nome']:<12} {s['cognome']:<15} voto: {s['voto']}")

    # ── STUDENTI A RISCHIO  (student_stats.studenti_a_rischio) ───────────────────
    rischio = studenti_a_rischio(
        studenti=validi,
        soglia_voto=float(config.get("soglia_voto", 6.0)),
        soglia_assenze=int(config.get("soglia_assenze", 15)),
    )
    righe.append("")
    righe.append("-" * 60)
    righe.append("  STUDENTI A RISCHIO")
    righe.append("-" * 60)
    righe.append(f"  Media insufficiente (<{config.get('soglia_voto', 6.0)}):  {len(rischio['media_insufficiente'])} studenti")
    righe.append(f"  Troppe assenze (>{config.get('soglia_assenze', 15)}):       {len(rischio['troppe_assenze'])} studenti")
    righe.append(f"  Entrambi i problemi:          {len(rischio['entrambi'])} studenti")
    if rischio["entrambi"]:
        righe.append("")
        righe.append("  Dettaglio (entrambi i problemi):")
        for s in rischio["entrambi"]:
            righe.append(
                f"    ⚠  {s['nome']:<12} {s['cognome']:<15} "
                f"media: {s['media']:<5}  assenze: {s['assenze']}"
            )

    # ── TOP 5 STUDENTI  (student_stats.classifica_studenti → "migliori") ─────────
    righe.append("")
    righe.append("-" * 60)
    righe.append("  TOP 5 STUDENTI")
    righe.append("-" * 60)
    for i, s in enumerate(top5, 1):
        righe.append(
            f"  {i}. {s['nome']:<12} {s['cognome']:<15} "
            f"media: {s.get('media', s.get('media_personale', '—'))}  assenze: {s['assenze']}"
        )

    # ── ESITI FINALI  (da student_logic via cmd_logica) ──────────────────────
    ha_esito = validi and "esito" in validi[0]
    if ha_esito:
        n_promossi = sum(1 for s in validi if s["esito"] == "Promosso")
        n_bocciati = len(validi) - n_promossi
        righe.append("")
        righe.append("-" * 60)
        righe.append("  ESITI FINALI")
        righe.append("-" * 60)
        righe.append(f"  Promossi:   {n_promossi}")
        righe.append(f"  Bocciati:   {n_bocciati}")
        bocciati = [s for s in validi if s["esito"] == "Bocciato"]
        if bocciati:
            righe.append("")
            righe.append("  Dettaglio bocciati:")
            for s in bocciati:
                motivi_str = "; ".join(s.get("motivi_esito", []))
                righe.append(
                    f"    • {s['nome']:<12} {s['cognome']:<15} "
                    f"voto finale: {s.get('voto_finale', '—')}  "
                    f"debiti: {len(s.get('debiti', []))}  "
                    f"[{motivi_str}]"
                )

    # ── FASCE DI RENDIMENTO  (da student_logic via cmd_logica) ───────────────
    if fasce:
        righe.append("")
        righe.append("-" * 60)
        righe.append("  FASCE DI RENDIMENTO")
        righe.append("-" * 60)
        for nome_fascia, lista_fascia in fasce.items():
            righe.append(f"  {nome_fascia:<15} {len(lista_fascia):>3} studenti")
            for s in lista_fascia:
                debiti_str = ", ".join(s.get("debiti", [])) or "nessuno"
                righe.append(
                    f"      {s['nome']:<12} {s['cognome']:<15} "
                    f"media: {s.get('media_personale', '—'):<6} "
                    f"voto: {s.get('voto_finale', '—'):<4} "
                    f"debiti: {debiti_str}"
                )

    # ── CLASSIFICA PER MEDIA  (da student_logic → bubble_sort) ──────────────
    if classifica_media:
        righe.append("")
        righe.append("-" * 60)
        righe.append("  CLASSIFICA PER MEDIA (bubble sort)")
        righe.append("-" * 60)
        for pos, s in enumerate(classifica_media, 1):
            righe.append(
                f"  {pos:>3}. {s['nome']:<12} {s['cognome']:<15} "
                f"media: {s.get('media_personale', '—'):<6} "
                f"esito: {s.get('esito', '—')}"
            )

    # ── CLASSIFICA PER ASSENZE  (da student_logic → quick_sort) ─────────────
    if classifica_assenze:
        righe.append("")
        righe.append("-" * 60)
        righe.append("  CLASSIFICA PER ASSENZE (quick sort, crescente)")
        righe.append("-" * 60)
        for pos, s in enumerate(classifica_assenze, 1):
            righe.append(
                f"  {pos:>3}. {s['nome']:<12} {s['cognome']:<15} "
                f"assenze: {s['assenze']:<4} "
                f"media: {s.get('media_personale', '—')}"
            )

    # STEP 17 — GENERAZIONE GRAFICI (integrato nel report)
    if genera_grafici and validi and ha_esito:
        print("\n" + "─" * 60)
        print("  Generazione grafici in corso...")
        print("─" * 60)
        try:
            genera_tutti_i_grafici(validi, config)
            print("[Step 17] ✅ Grafici generati con successo in report/charts/")
        except Exception as e:
            print(f"[Step 17] ⚠  Avvertimento: impossibile generare grafici: {e}")

    righe.append("")
    righe.append("=" * 60)

    with open(percorso, "w", encoding="utf-8") as f:
        f.write("\n".join(righe))



    # 1. Crea un'istanza del documento
    # 2. Aggiungi un titolo (livello 0 è il più grande, poi 1, 2...)
    # 3. Aggiungi un paragrafo normale
    # 4. Aggiungi testo formattato (grassetto) all'interno dello stesso paragrafo
    # 5. Salva il documento



    # doc = Document()
    # doc.add_heading('Titolo del Documento', level=0)
    # paragrafo = doc.add_paragraph('Questo è un semplice paragrafo scritto con Python.')
    # paragrafo.add_run(' Questo testo è in grassetto.').bold = True
    # doc.save('documento.docx')



    print(f"[Step 8] Report salvato in: {percorso}")
    return percorso
'''





def backup_csv(percorso_csv : Path) -> Path:            # STEP 9 - Backup automatico
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    directory = Path("data","backup")          #("PROGETTO-Student Analytics Pipeline","data","backup")
    destinazione = directory / f"backup_studenti_{timestamp}.csv"
    shutil.copy2(percorso_csv, destinazione)
    print(f"[Step 9] Backup salvato in {destinazione}")
    return destinazione
##################################################################################################################################################
##################################################################################################################################################
def cmd_gendb():           # STEP 10 — Gestione CLI con sys.argv
    database_main()         # STEP 1 - caricamento configdb.json e generazione db
    print("[Step Alpha] Database generato con successo.")
##################################################################################################################################################
def cmd_generate(config):           # STEP 10 — Gestione CLI con sys.argv
    config_db = carica_configdb()       # STEP 1 - caricamento configdb.json
    studenti = leggi_studenti_da_db(config_db["NOME_DATABASE"] + ".db")       # STEP 2 - Generazione studenti tramite database
    lista_csv = salva_su_csv(studenti= studenti ,config= config)       # STEP 3
    aggiungi_errori(lista_csv, config)         # STEP 2B
    backup_csv(lista_csv)         # STEP 9
    return lista_csv
##################################################################################################################################################
def cmd_validate(config):           # STEP 10 — Gestione CLI con sys.argv
    directory = Path("data/input")
    file_csv = sorted(directory.glob("studenti_*.csv"))
    if not file_csv:
        print("[ERRORE] Nessun file CSV trovato in data/input/. Esegui prima 'generate'.")
        sys.exit(1)
    percorso_csv = file_csv[-1]         # solo per stampare in ordine alfabetico
    print(f"[Step 4] Lettura file: {percorso_csv}")

    validi, scartati, _ = valida_studenti(percorso_csv, config)       # STEP 4
    salva_json_validi(validi)            # STEP 5
    salva_json_scartati(scartati)            # STEP 5
    return validi, scartati
##################################################################################################################################################
def cmd_logica(config):           # STEP 10 — Gestione CLI con sys.argv
    """Esegue la pipeline di logica (STEP 11-16) su studenti_validi.json."""
    directory = Path("data/output")
    json_path = directory / "studenti_validi.json"
    if not json_path.exists():
        print("[ERRORE] studenti_validi.json non trovato. Esegui prima 'validate'.")
        sys.exit(1)
    with open(json_path, "r", encoding="utf-8") as f:
        validilog = json.load(f)

    calcola_media_e_voto_finale(validilog, config)              # STEP 11
    assegna_debiti_formativi(validilog, config)                 # STEP 12
    determina_esito(validilog, config)                          # STEP 13
    fasce = raggruppa_per_fascia_rendimento(validilog, config)  # STEP 14

    classifica_media   = bubble_sort_studenti(validilog, "media_personale")               # STEP 15
    classifica_assenze = quick_sort_studenti(validilog, "assenze", decrescente=False)     # STEP 16

    return validilog, fasce, classifica_media, classifica_assenze
##################################################################################################################################################
def cmd_charts(config):
    """Genera solo i grafici (senza rigenerare il report)."""
    directory = Path("data/output")
    json_path = directory / "studenti_validi.json"
    
    # Carica studenti
    with open(json_path, "r", encoding="utf-8") as f:
        validi = json.load(f)

    # Se non hanno esito, esegui logica
    ha_esito = validi and "esito" in validi[0]
    if not ha_esito:
        print("[AVVISO] I dati non hanno ancora l'esito calcolato. Eseguo logica...")
        validi, _, _, _ = cmd_logica(config)

    # Genera grafici
    genera_tutti_i_grafici(validi, config)
    print("\n[Step 17] Grafici generati con successo in report/charts/")
##################################################################################################################################################
def cmd_report(config):           # STEP 10 — Gestione CLI con sys.argv
    directory = Path("data/output")

    json_path = directory / "studenti_validi.json"
    if not json_path.exists():
        print("[ERRORE] studenti_validi.json non trovato. Esegui prima 'validate'.")
        sys.exit(1)
    with open(json_path, "r", encoding="utf-8") as f:
        validimedia = json.load(f)

    scartati_path = directory / "studenti_scartati.json"
    scartati = []
    if scartati_path.exists():
        with open(scartati_path, "r", encoding="utf-8") as f:
            scartati = json.load(f)

    statistica_per_materia = media_per_materia(validimedia)         # STEP 6
    classifica = classifica_studenti(validimedia, top_n=5)["migliori"]                   # STEP 7

    validilog, fasce, classifica_media, classifica_assenze = cmd_logica(config)   # STEP 11-16

    genera_report_docx(config, validilog, scartati, statistica_per_materia, classifica,    # STEP 8
                    fasce=fasce, classifica_media=classifica_media,
                    classifica_assenze=classifica_assenze)
##################################################################################################################################################
def cmd_aggregate_generate_validate(config):           # STEP 10 — Gestione CLI con sys.argv
    """Esegue generate + validate in sequenza (senza report)."""
    print("\n FASE 1 — Generazione dati")
    cmd_generate(config)
    print("\n FASE 2 — Validazione e conversione")
    cmd_validate(config)
##################################################################################################################################################
def cmd_all(config):           # STEP 10 — Gestione CLI con sys.argv
    print("\n FASE Alpha — Generazione database")
    cmd_gendb()         # STEP 1 - caricamento configdb.json e generazione db

    print("\n FASE 1 — Generazione dati")
    cmd_generate(config)

    print("\n FASE 2 — Validazione e conversione")
    validimedia, scartati = cmd_validate(config)

    print("\n FASE 3 — Statistiche e report")
    validilog, fasce, classifica_media, classifica_assenze = cmd_logica(config)
    statistica_per_materia = media_per_materia(validimedia)         # STEP 6
    classifica = classifica_studenti(validimedia, top_n=5)["migliori"]                   # STEP 7

    genera_report_docx(config, validilog, scartati, statistica_per_materia, classifica,      # STEP 8
                fasce=fasce, classifica_media=classifica_media,
                classifica_assenze=classifica_assenze)

    print("\n Pipeline completata con successo.")
##################################################################################################################################################
def mosta_aiuto():           # STEP 10 — Gestione CLI con sys.argv
    print("""
        Uso:  python progetto.py <comando>
        
        E' necessario avere la libreria DOCX per far girare il programma, 

                                -----------
        LANCIA IL COMANDO :     | getDOCX |       per avviare il processo di installazione
                                -----------
        
        Comandi disponibili:
            gendb                   - Genera il database SQLite con studenti che arrivano dal percorso selezionato in input
            generate                - Genera in base a una lista di studenti inserita in input un file CSV in data/input/
            validate                - Valida il CSV e produce JSON in data/output/
            gv / generate-validate  - Esegue la generazione e la validazione in sequenza (senza report)
            logica                  - Calcola media, debiti, esito e classifiche su studenti_validi.json
            report                  - Calcola statistiche su un file CSV esistente e genera il report in report/
            all                     - Esegue l'intera pipeline dall'inizio alla fine
        """)
##################################################################################################################################################
##################################################################################################################################################
# def main():         # STEP CLI - GESTIONE INPUT
if __name__ == "__main__":
    crea_cartelle()            # STEP 0 
    config = carica_config()       # STEP 1
    
    if len(sys.argv) <2:
        mosta_aiuto()
        sys.exit(0)

    # comando = input().lower()
    comando = sys.argv[1].lower()

    if comando == "":
        mosta_aiuto()
    elif comando == "getdocx":
        install_procedure()
    elif comando == "gendb":
        cmd_gendb()
    elif comando == "generate":
        cmd_generate(config)
    elif comando == "validate":
        cmd_validate(config)
    elif comando == "logica":
        cmd_logica(config)    
    elif comando == "report":
        cmd_report(config)
    elif comando in ("generate-validate", "gv"):
        cmd_aggregate_generate_validate(config)
    elif comando == "all":
        cmd_all(config)
    else:
        print(f"[ERRORE] Comando sconosciuto: '{comando}'")
        mosta_aiuto()
        sys.exit(1)
